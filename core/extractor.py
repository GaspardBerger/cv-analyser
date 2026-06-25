#!/usr/bin/env python3
"""Tekstextractie uit PDF- en DOCX-bestanden voor CV-analyse.

Strategie voor PDF's:
1. pdfplumber  — snel, werkt voor tekst-gebaseerde PDF's
2. PyMuPDF     — tweede poging, beter bij sommige PDF-types
3. Claude OCR  — fallback voor gescande/afbeelding-PDF's (via Vision API)
"""

from pathlib import Path


def extraheer_tekst(bestandspad: str) -> tuple[str, str | None]:
    """
    Extraheer leesbare tekst uit een PDF of DOCX bestand.

    Geeft terug: (tekst, melding)
    - tekst:    geëxtraheerde CV-tekst (leeg bij mislukking)
    - melding:  None bij succes; foutmelding als tekst leeg is;
                informatieve melding als OCR werd gebruikt (tekst is dan WEL gevuld)
    """
    pad = Path(bestandspad)
    suffix = pad.suffix.lower()

    if suffix == ".pdf":
        return _extraheer_pdf(bestandspad)
    elif suffix in (".docx", ".doc"):
        return _extraheer_docx(bestandspad)
    else:
        return "", f"Bestandsformaat '{suffix}' wordt niet ondersteund. Gebruik PDF of DOCX."


# ── PDF ──────────────────────────────────────────────────────────────────────

def _extraheer_pdf(pad: str) -> tuple[str, str | None]:
    # Stap 1: pdfplumber
    tekst = _probeer_pdfplumber(pad)
    if tekst:
        return tekst, None

    # Stap 2: PyMuPDF (soms beter voor hybride PDF's)
    tekst = _probeer_pymupdf_tekst(pad)
    if tekst:
        return tekst, None

    # Stap 3: OCR via Claude Vision (gescande PDF's)
    return _ocr_via_claude_vision(pad)


def _probeer_pdfplumber(pad: str) -> str:
    try:
        import pdfplumber
    except ImportError:
        return ""
    try:
        tekst_delen = []
        with pdfplumber.open(pad) as pdf:
            for pagina in pdf.pages:
                inhoud = pagina.extract_text()
                if inhoud:
                    tekst_delen.append(inhoud)
        tekst = "\n".join(tekst_delen).strip()
        return tekst if len(tekst) >= 50 else ""
    except Exception:
        return ""


def _probeer_pymupdf_tekst(pad: str) -> str:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return ""
    try:
        doc = fitz.open(pad)
        tekst_delen = []
        for pagina in doc:
            inhoud = pagina.get_text()
            if inhoud and inhoud.strip():
                tekst_delen.append(inhoud.strip())
        tekst = "\n".join(tekst_delen).strip()
        return tekst if len(tekst) >= 50 else ""
    except Exception:
        return ""


def _ocr_via_claude_vision(pad: str) -> tuple[str, str | None]:
    """Render PDF-pagina's als afbeeldingen en stuur ze naar Claude Vision voor OCR."""
    import base64
    import os

    try:
        import fitz  # PyMuPDF
    except ImportError:
        return "", (
            "Dit lijkt een gescande PDF. Installeer PyMuPDF om gescande CV's te lezen: "
            "'pip install PyMuPDF'."
        )

    try:
        import anthropic
    except ImportError:
        return "", "anthropic is niet geïnstalleerd."

    api_sleutel = os.environ.get("ANTHROPIC_API_KEY")
    if not api_sleutel:
        return "", "API-sleutel niet beschikbaar voor OCR."

    try:
        doc = fitz.open(pad)
    except Exception as e:
        return "", f"Kon PDF niet openen: {e}"

    # Een CV is maximaal 2–3 pagina's; we verwerken max. 5 pagina's
    max_paginas = min(len(doc), 5)
    client = anthropic.Anthropic(api_key=api_sleutel)
    tekst_delen = []

    for i in range(max_paginas):
        pagina = doc[i]
        # 150 DPI: goede kwaliteit zonder te grote bestanden
        mat = fitz.Matrix(150 / 72, 150 / 72)
        pix = pagina.get_pixmap(matrix=mat)
        img_b64 = base64.standard_b64encode(pix.tobytes("png")).decode()

        try:
            bericht = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=2000,
                messages=[{
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": "image/png",
                                "data": img_b64,
                            },
                        },
                        {
                            "type": "text",
                            "text": (
                                "Extraheer alle tekst uit dit CV-beeld zo volledig en nauwkeurig mogelijk. "
                                "Behoud de structuur (secties, koppen). "
                                "Geef alleen de tekst terug, zonder commentaar of uitleg."
                            ),
                        },
                    ],
                }],
            )
            pagina_tekst = bericht.content[0].text.strip()
            if pagina_tekst:
                tekst_delen.append(pagina_tekst)
        except Exception:
            continue

    tekst = "\n\n".join(tekst_delen).strip()

    if not tekst or len(tekst) < 50:
        return "", (
            "Er kon geen tekst worden uitgelezen uit dit document, ook niet via OCR. "
            "Probeer het CV te exporteren als PDF vanuit Word of Google Docs."
        )

    # Tekst gevonden via OCR — informatie voor de gebruiker (niet-blokkerend)
    ocr_melding = "ocr_gebruikt"
    return tekst, ocr_melding


# ── DOCX ─────────────────────────────────────────────────────────────────────

def _extraheer_docx(pad: str) -> tuple[str, str | None]:
    try:
        from docx import Document
    except ImportError:
        return "", "python-docx is niet geïnstalleerd. Voer 'pip install python-docx' uit."

    try:
        doc = Document(pad)
    except Exception as fout:
        fout_str = str(fout).lower()
        if "encrypted" in fout_str or "password" in fout_str or "corrupt" in fout_str:
            return "", (
                "Dit Word-bestand is beveiligd met een wachtwoord of beschadigd. "
                "Sla het op als een nieuw bestand zonder wachtwoordbeveiliging."
            )
        return "", f"Kon het Word-bestand niet openen: {fout}"

    alineas = [p.text for p in doc.paragraphs if p.text.strip()]

    # Tekst in tabellen (veel CV's gebruiken tabellen voor lay-out)
    for tabel in doc.tables:
        for rij in tabel.rows:
            for cel in rij.cells:
                cel_tekst = cel.text.strip()
                if cel_tekst and cel_tekst not in alineas:
                    alineas.append(cel_tekst)

    tekst = "\n".join(alineas).strip()

    if not tekst or len(tekst) < 50:
        return "", "Het Word-bestand lijkt leeg of bevat alleen afbeeldingen."

    return tekst, None
